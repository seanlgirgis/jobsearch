#!/usr/bin/env python3
"""
scripts/resume_generation.py
Renders tailored resume from intermediate JSON to MD preview and DOCX.
Supports trimming for 1-2 page resumes (up to 5 recent jobs + summary block).
Now includes --all to generate both trimmed and untrimmed versions at once.

Dependencies: pip install python-docx

Usage:
    python scripts/resume_generation.py --uuid <uuid> --version v1 [--trim | --all]
"""
import argparse
from pathlib import Path
import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

JOB_ROOT = Path("data/jobs")
MASTER_ROOT = Path("data/master")


def load_intermediate(uuid: str, version: str) -> dict:
    job_dir = JOB_ROOT / uuid / "generated"
    json_path = job_dir / f"resume_intermediate_{version}.json"
    if not json_path.is_file():
        raise FileNotFoundError(f"Missing intermediate JSON: {json_path}")
    with open(json_path, "r", encoding="utf-8") as f:
        return json.load(f)


def load_exclusions() -> set:
    career_path = MASTER_ROOT / "master_career_data.json"
    if not career_path.is_file():
        print("Warning: master_career_data.json not found — no exclusions applied")
        return set()
    with open(career_path, "r", encoding="utf-8") as f:
        master = json.load(f)
    excluded = {exp.get("company", "") for exp in master.get("experience", []) if exp.get("exclude_from_resume", False)}
    if excluded:
        print(f"Applying exclusions for companies: {excluded}")
    return excluded


def clean_url_for_display(url: str) -> str:
    """Shorten URL for display (remove protocol, www, trailing slash)"""
    if not url:
        return ""
    cleaned = url.replace("https://", "").replace("http://", "").replace("www.", "").rstrip("/")
    return cleaned or url


def render_markdown(tailored: dict, trim: bool = False) -> str:
    md = "# Tailored Resume Preview\n\n"

    # Compact header
    personal = tailored.get("personal", {})
    name = personal.get("full_name", "Name Missing")
    title = personal.get("preferred_title", "")
    phone = personal.get("phone", "")
    email = personal.get("email", "")
    linkedin = personal.get("linkedin", "")
    github = personal.get("github", "")
    website = personal.get("personal_website", "")

    md += f"**{name}**  \n"
    if title:
        md += f"{title}  \n"
    md += f"{phone} | {email}  \n"

    links = []
    if linkedin:
        links.append(f"[{clean_url_for_display(linkedin)}]({linkedin})")
    if github:
        links.append(f"[{clean_url_for_display(github)}]({github})")
    if website:
        links.append(f"[{clean_url_for_display(website)}]({website})")
    if links:
        md += " | ".join(links) + "  \n\n"

    # Summary
    md += "## Professional Summary\n"
    md += tailored.get("summary", "Summary missing") + "\n\n"

    # Experience (with exclusions applied)
    excluded_companies = load_exclusions()
    experiences = [exp for exp in tailored.get("experience", []) if exp.get("company") not in excluded_companies]
    md += "## Professional Experience\n"
    display_exps = experiences[:5] if trim else experiences
    for exp in display_exps:
        md += f"### {exp.get('title', 'Title')} at {exp.get('company', 'Company')}\n"
        md += f"{exp.get('start_date', '?')} – {exp.get('end_date', '?')}\n\n"
        for bullet in exp.get("bullets", [])[:7]:
            md += f"- {bullet}\n"
        md += "\n"

    # NO "Additional Experience" in trimmed mode
    if not trim and len(experiences) > 5:
        md += "**Additional Experience** (earlier roles):  \n"
        md += ", ".join(
            f"{e.get('title','?')} at {e.get('company','?')} ({e.get('start_date','')}-{e.get('end_date','')})"
            for e in experiences[5:]
        ) + "\n\n"

    # Education
    md += "## Education\n"
    for edu in tailored.get("education", []):
        degree = edu.get("degree", "Degree")
        institution = edu.get("institution", "")
        location = edu.get("location", "")
        dates = edu.get("dates", "")

        md += f"- **{degree}**  \n"
        md += f"  {institution}"
        if location:
            md += f", {location}"
        if dates and dates.lower() not in ["not specified", "(not specified)", "", "unknown"]:
            md += f" ({dates})"
        md += "\n"
    md += "\n"

    # Skills
    md += "## Skills\n"
    skills = tailored.get("skills", [])
    skill_names = [s.get("name", s) if isinstance(s, dict) else s for s in skills]
    md += ", ".join(filter(None, skill_names)) + "\n\n"

    # Flagship Projects (only if present)
    projects = tailored.get("projects", []) or tailored.get("flagship_projects", [])
    if projects:
        md += "## Flagship Projects\n"
        for proj in projects:
            md += f"### {proj.get('name', 'Project Name')}\n"
            md += proj.get("description", "") + "\n"
            if proj.get("technologies"):
                md += f"Technologies: {', '.join(proj['technologies'])}\n"
            repo = proj.get("repo", "")
            if repo:
                repo_clean = repo.split('https://')[-1].strip() if 'https://' in repo else repo
                repo_display = clean_url_for_display(repo_clean or repo)
                if repo_display:
                    md += f"Repo: [{repo_display}]({repo})\n"
            md += "\n"

    return md


def add_hyperlink(paragraph, url: str, text: str):
    """Add a real clickable hyperlink to a docx paragraph"""
    if not url or not text:
        paragraph.add_run(text)  # fallback to plain text
        return
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    new_run.append(rPr)
    text_el = OxmlElement('w:t')
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def render_docx(tailored: dict, out_path: Path, trim: bool = False):
    doc = Document()

    personal = tailored.get("personal", {})
    name = personal.get("full_name", "Name Missing")
    title = personal.get("preferred_title", "")
    phone = personal.get("phone", "(phone missing)")
    email = personal.get("email", "(email missing)")
    linkedin = personal.get("linkedin", "")
    github = personal.get("github", "")
    website = personal.get("personal_website", "")

    # Compact centered header – same style for trimmed & full
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(16)

    if title:
        p.add_run("\n")
        run = p.add_run(title)
        run.font.size = Pt(12)

    p.add_run("\n")
    p.add_run(phone)
    p.add_run("  |  ")
    p.add_run(email)

    links = [(linkedin, "LinkedIn"), (github, "GitHub"), (website, "Portfolio")]
    valid_links = [(url, label) for url, label in links if url and "http" in url]

    if valid_links:
        p.add_run("\n")
        for i, (url, label) in enumerate(valid_links):
            if i > 0:
                p.add_run("  |  ")
            add_hyperlink(p, url, label)

    # Summary
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(tailored.get("summary", ""))

    # Experience (with exclusions applied)
    excluded_companies = load_exclusions()
    experiences = [exp for exp in tailored.get("experience", []) if exp.get("company") not in excluded_companies]
    doc.add_heading("Professional Experience", level=1)
    display_exps = experiences[:5] if trim else experiences

    for exp in display_exps:
        p = doc.add_paragraph()
        run = p.add_run(f"{exp.get('title', 'Title')} at {exp.get('company', 'Company')}")
        run.bold = True
        p.add_run(f"\n{exp.get('start_date', '?')} – {exp.get('end_date', '?')}")
        for bullet in exp.get("bullets", [])[:7]:
            doc.add_paragraph(bullet, style="List Bullet")

    # NO "Additional Experience" in trimmed mode
    if not trim and len(experiences) > 5:
        p = doc.add_paragraph()
        p.add_run("**Additional Experience** (earlier roles): " + ", ".join(
            f"{e.get('title','?')} at {e.get('company','?')} ({e.get('start_date','')}-{e.get('end_date','')})"
            for e in experiences[5:]
        ))

    # Education
    doc.add_heading("Education", level=1)
    for edu in tailored.get("education", []):
        p = doc.add_paragraph()
        run = p.add_run(edu.get("degree", "Degree"))
        run.bold = True

        line = f"\n{edu.get('institution', '')}"
        if edu.get('location'):
            line += f", {edu.get('location')}"
        dates = edu.get("dates", "")
        if dates and dates.lower() not in ["not specified", "(not specified)", "", "unknown"]:
            line += f" ({dates})"

        p.add_run(line)

    # Skills
    doc.add_heading("Skills", level=1)
    skills = tailored.get("skills", [])
    skill_str = ", ".join(s.get("name", s) if isinstance(s, dict) else s for s in skills)
    doc.add_paragraph(skill_str)

    # Flagship Projects (only if present)
    projects = tailored.get("projects", []) or tailored.get("flagship_projects", [])
    if projects:
        doc.add_heading("Flagship Projects", level=1)
        for proj in projects:
            p = doc.add_paragraph()
            run = p.add_run(proj.get("name", "Project"))
            run.bold = True
            p.add_run(f"\n{proj.get('description', '')}")
            if proj.get("technologies"):
                p.add_run(f"\nTechnologies: {', '.join(proj['technologies'])}")
            repo = proj.get("repo", "")
            if repo:
                repo_clean = repo.split('https://')[-1].strip() if 'https://' in repo else repo
                repo_display = clean_url_for_display(repo_clean or repo)
                if repo_display:
                    p.add_run(f"\nRepo: ")
                    add_hyperlink(p, repo, repo_display)

    # Formatting
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial"
            if paragraph.style.name.startswith("Heading"):
                run.font.size = Pt(13)
            else:
                run.font.size = Pt(11)

    doc.save(out_path)


def generate_files(tailored, uuid, version, trim_mode):
    suffix = "_trimmed" if trim_mode else ""
    md_path = JOB_ROOT / uuid / "generated" / f"resume_preview_{version}{suffix}.md"
    docx_path = JOB_ROOT / uuid / "generated" / f"resume_{version}{suffix}.docx"

    with open(md_path, "w", encoding="utf-8") as f:
        f.write(render_markdown(tailored, trim=trim_mode))
    print(f"Markdown saved → {md_path}")

    render_docx(tailored, docx_path, trim=trim_mode)
    print(f"DOCX saved → {docx_path}")


def main():
    parser = argparse.ArgumentParser(description="Render resume from intermediate JSON to MD & DOCX")
    parser.add_argument("--uuid", required=True, help="Job UUID")
    parser.add_argument("--version", default="v1", help="Version tag")
    group = parser.add_mutually_exclusive_group()
    group.add_argument("--trim", action="store_true", help="Generate trimmed version only (5 recent jobs)")
    group.add_argument("--all", action="store_true", help="Generate both trimmed and full versions")
    args = parser.parse_args()

    job_dir = JOB_ROOT / args.uuid / "generated"
    job_dir.mkdir(parents=True, exist_ok=True)

    tailored = load_intermediate(args.uuid, args.version)

    if args.all:
        print("Generating both trimmed and full versions...")
        generate_files(tailored, args.uuid, args.version, trim_mode=True)
        generate_files(tailored, args.uuid, args.version, trim_mode=False)
    else:
        trim_mode = args.trim
        mode_name = "trimmed" if trim_mode else "full"
        print(f"Generating {mode_name} version...")
        generate_files(tailored, args.uuid, args.version, trim_mode=trim_mode)

    print("\nDone.")


if __name__ == "__main__":
    main()