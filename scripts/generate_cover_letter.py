#!/usr/bin/env python3
"""
scripts/generate_cover_letter.py

FIXED VERSION:
- Dynamic company + role
- POSITIONING_ANGLE support
- Validation layer added
- Safer structure rules
"""

import argparse
import json
import yaml
from pathlib import Path
from docx import Document
from docx.shared import Pt

from src.ai.grok_client import GrokClient

JOB_ROOT = Path("data/jobs")
MASTER_ROOT = Path("data/master")


COVER_PROMPT = """
You are an expert cover letter writer for senior data engineering roles.

CRITICAL RULES — VIOLATE ANY AND OUTPUT IS INVALID:
- Structure:
  - Opening paragraph
  - 1–2 body paragraphs
  - Closing paragraph + sign-off
- Length: 300–450 words
- Tone: First-person, professional, confident, concise
- No fabrication: ONLY use facts from MASTER CAREER DATA and JOB DATA
- Do NOT output JSON, headings, or explanations

POSITIONING_ANGLE:
{positioning_angle}

PERSONALIZATION:
- Company: {company_name}
- Role: {job_title}

CONTENT REQUIREMENTS:
- Opening MUST reflect POSITIONING_ANGLE clearly
- Highlight 3–5 relevant achievements (quantified if possible)
- Emphasize AWS, Python, ETL, and ML-enabled pipelines where applicable
- ML should be positioned as enablement (NOT deep learning)

MASTER CAREER DATA:
{master_career_data}

JOB DATA:
{job_data}

Start with:
Dear Hiring Manager,

End with:
Best regards,
Sean Luka Girgis
214-315-2190
seanlgirgis@gmail.com
https://www.linkedin.com/in/sean-girgis-43bb1b5/

Output ONLY the plain text cover letter.
"""


def load_master():
    with open(MASTER_ROOT / "master_career_data.json", "r") as f:
        return json.load(f)


def load_job(uuid_str: str):
    job_dir = JOB_ROOT / uuid_str
    with open(job_dir / "tailored" / "tailored_data_v1.yaml", "r") as f:
        return yaml.safe_load(f)


def validate_letter(text: str):
    if "Dear Hiring Manager" not in text:
        raise ValueError("Missing greeting")
    if "Best regards" not in text:
        raise ValueError("Missing sign-off")
    if len(text.split()) < 250:
        raise ValueError("Letter too short")
    return text


def generate_tailored_letter(uuid_str: str, positioning_angle: str, llm: str = "grok"):
    master_career = load_master()
    job_data = load_job(uuid_str)

    prompt = COVER_PROMPT.format(
        positioning_angle=positioning_angle,
        company_name=job_data.get("company_name", ""),
        job_title=job_data.get("job_title", ""),
        master_career_data=json.dumps(master_career, indent=2),
        job_data=yaml.safe_dump(job_data, sort_keys=False)
    )

    client = GrokClient()
    response = client.query(
        prompt,
        model="grok-beta",
        temperature=0.1,
        max_tokens=1000
    )

    return validate_letter(response.strip())


def render_docx(letter_text: str, out_path: Path):
    doc = Document()
    for para in letter_text.split("\n\n"):
        doc.add_paragraph(para)
    for p in doc.paragraphs:
        p.style.font.name = "Arial"
        p.style.font.size = Pt(11)
    doc.save(out_path)


def main():
    parser = argparse.ArgumentParser(description="Generate tailored cover letter")
    parser.add_argument("--uuid", required=True)
    parser.add_argument("--angle", required=True)
    args = parser.parse_args()

    job_dir = JOB_ROOT / args.uuid / "generated"
    job_dir.mkdir(parents=True, exist_ok=True)

    letter_text = generate_tailored_letter(args.uuid, args.angle)

    txt_path = job_dir / "cover_letter_v1.txt"
    with open(txt_path, "w") as f:
        f.write(letter_text)

    docx_path = job_dir / "cover_letter_v1.docx"
    render_docx(letter_text, docx_path)

    print(f"Generated: {txt_path}, {docx_path}")


if __name__ == "__main__":
    main()
