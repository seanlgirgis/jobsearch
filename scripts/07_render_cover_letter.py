#!/usr/bin/env python3
"""
scripts/07_render_cover_letter.py

Phase 7: Render cover letter from intermediate JSON to DOCX and MD.
Supports agency/generic vs enterprise/personalized (from JSON).

Usage:
    python scripts/07_render_cover_letter.py --uuid <uuid-or-prefix> --version v1
"""

import argparse
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

JOB_ROOT = Path("data/jobs")

def resolve_job_folder(uuid_str: str) -> Path:
    """Smart resolution consistent with phases 02-06"""
    uuid_str = uuid_str.strip()
    job_dir = JOB_ROOT / uuid_str

    if job_dir.is_dir():
        return job_dir

    candidates = list(JOB_ROOT.glob(f"*_{uuid_str}*"))
    if not candidates:
        short = uuid_str[:8]
        candidates = list(JOB_ROOT.glob(f"*_{short}*"))

    if len(candidates) == 1:
        print(f"Resolved UUID/prefix → using folder: {candidates[0].name}")
        return candidates[0]
    elif len(candidates) > 1:
        print("Ambiguous UUID — multiple matches:")
        for c in candidates:
            print(f"  - {c.name}")
        raise ValueError(f"Ambiguous UUID: {uuid_str}")
    else:
        raise ValueError(
            f"Job folder not found for '{uuid_str}'\n"
            f"Expected: data/jobs/{uuid_str} or data/jobs/*_{uuid_str}*"
        )


def load_cover_intermediate(job_dir: Path, version: str) -> dict:
    """Load latest cover_intermediate_*.json"""
    generated_dir = job_dir / "generated"
    if not generated_dir.is_dir():
        raise FileNotFoundError(f"No generated/ folder: {generated_dir}")

    # Try exact version first
    exact_path = generated_dir / f"cover_intermediate_{version}.json"
    if exact_path.is_file():
        print(f"Using exact cover intermediate: {exact_path.name}")
        with open(exact_path, "r", encoding="utf-8") as f:
            return json.load(f)

    # Fallback to latest
    files = sorted(generated_dir.glob("cover_intermediate_*.json"), reverse=True)
    if not files:
        raise FileNotFoundError("No cover_intermediate_*.json found")
    latest = files[0]
    print(f"Exact version not found — using latest: {latest.name}")
    with open(latest, "r", encoding="utf-8") as f:
        return json.load(f)


def render_markdown(cover: dict) -> str:
    md = "# Tailored Cover Letter Preview\n\n"

    # Header
    header = cover.get("header", {})
    md += f"**{header.get('name', 'Your Name')}**  \n"
    md += f"{header.get('address', '')}  \n"
    md += f"{header.get('phone', '')} | {header.get('email', '')}  \n\n"
    md += f"{header.get('date', 'Date')}  \n\n"
    md += f"{header.get('employer_address', 'Hiring Manager')}\n\n"

    # Salutation
    md += f"{cover.get('salutation', 'Dear Hiring Manager,')}\n\n"

    # Intro
    md += f"{cover.get('intro', '')}\n\n"

    # Body
    for para in cover.get("body", []):
        md += f"{para}\n\n"

    # Conclusion
    md += f"{cover.get('conclusion', '')}\n\n"

    # Sign-off
    md += f"{cover.get('sign_off', 'Sincerely,\\nYour Name')}\n"

    return md


def render_docx(cover: dict, out_path: Path):
    doc = Document()

    # Header (your info left-aligned)
    header_p = doc.add_paragraph()
    header = cover.get("header", {})
    header_p.add_run(f"{header.get('name', 'Your Name')}\n")
    header_p.add_run(f"{header.get('address', '')}\n")
    header_p.add_run(f"{header.get('phone', '')} | {header.get('email', '')}\n")
    header_p.add_run(f"{header.get('date', 'Date')}\n\n")

    # Employer address (right-aligned)
    employer_p = doc.add_paragraph()
    employer_p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    employer_p.add_run(f"{header.get('employer_address', 'Hiring Manager')}\n")

    # Salutation
    doc.add_paragraph(cover.get("salutation", "Dear Hiring Manager,"))

    # Intro paragraph
    doc.add_paragraph(cover.get("intro", ""))

    # Body paragraphs
    for para in cover.get("body", []):
        doc.add_paragraph(para)

    # Conclusion
    doc.add_paragraph(cover.get("conclusion", ""))

    # Sign-off
    sign_p = doc.add_paragraph(cover.get("sign_off", "Sincerely,\nYour Name"))
    sign_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Formatting
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)

    doc.save(out_path)


def main():
    parser = argparse.ArgumentParser(description="Phase 7: Render cover letter to DOCX/MD")
    parser.add_argument("--uuid", required=True, help="Job UUID or short prefix (e.g. cdb9a3fa)")
    parser.add_argument("--version", default="v1", help="Intermediate JSON version tag (e.g. v1)")
    args = parser.parse_args()

    job_folder = resolve_job_folder(args.uuid)
    print(f"Processing job: {job_folder.name}")

    cover = load_cover_intermediate(job_folder, args.version)

    # Markdown preview
    md_path = job_folder / "generated" / f"cover_preview_{args.version}.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(render_markdown(cover))
    print(f"Markdown saved → {md_path}")

    # DOCX letter
    docx_path = job_folder / "generated" / f"cover_letter_{args.version}.docx"
    render_docx(cover, docx_path)
    print(f"DOCX saved → {docx_path}")

    print("\nDone.")


if __name__ == "__main__":
    main()