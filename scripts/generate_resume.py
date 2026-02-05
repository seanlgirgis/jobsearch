"""
scripts/generate_resume.py
Generates a tailored resume in structured JSON format (via LLM or mock),
then renders it to Markdown preview and improved DOCX with styling & real hyperlinks.
Dependencies:
    pip install python-docx pyyaml
Usage examples:
    # Full generation (calls LLM)
    python -m scripts.generate_resume --uuid 96b16121-8608-405d-9553-af86fdbf939c --version v4
    # Re-render from latest JSON (no LLM call)
    python -m scripts.generate_resume --uuid 96b16121-8608-405d-9553-af86fdbf939c --version v7-custom --skip-llm
"""
import argparse
import json
import yaml
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# Try to import GrokClient — if missing → mock mode
try:
    from src.ai.grok_client import GrokClient
except ImportError:
    print("Warning: src.ai.grok_client not found → entering MOCK / MANUAL mode.")
    GrokClient = None

JOB_ROOT = Path("data/jobs")
MASTER_ROOT = Path("data/master")

# ────────────────────────────────────────────────
# Final prompt with proper escaping (Variation 3.1)
# ────────────────────────────────────────────────
RESUME_PROMPT = """
You are a rule-bound resume optimizer. Temperature=0 for determinism. Violate rules = invalid response.
CHECKLIST (MUST ADHERE 100%):
1. No truncation: Keep EVERY section/job/project/education from master.
2. Bullet count: >= master's per role (no reductions).
3. Length: Ensure full depth (1.5–2 pages equiv.).
4. No removals: Education/Skills/Flagships ALWAYS included.
5. Rephrase only: Match job keywords (e.g., "SQL", "data governance") naturally; quantify existing metrics.
6. Format: Action verb starts, no "I", ATS-safe.
7. Skills: Full master list, top 5–10 job-matched first.
8. Education: VERBATIM MEANS EXACT WORD-FOR-WORD COPY from master; no changes to degrees/institutions (e.g., if master says "Post-Graduate Diploma in Computer Science", output EXACTLY that).
9. Personal, Education, Flagship Projects: Copy fields verbatim from master_career_data where possible; do NOT change degree names, institutions, project names, or core facts (e.g., if master has 2 flagship projects, output EXACTLY 2 with rephrased descriptions only).
10. Experience roles: Only use companies/titles/dates exactly as they appear in master; do NOT invent, split, or add new roles (e.g., if master has 4 experience entries, output EXACTLY 4).
EXAMPLE FOR VERBATIM COPY:
Master education: [{{"degree": "Post-Graduate Diploma in Computer Science", "institution": "Humber College", "location": "Canada"}}]
Output education: [{{"degree": "Post-Graduate Diploma in Computer Science", "institution": "Humber College", "location": "Canada"}}] — NO CHANGES.
Master flagship_projects: 2 entries
Output flagship_projects: EXACTLY 2 entries, rephrased descriptions only.
MASTER CAREER DATA: {master_career_data}
MASTER SKILLS: {master_skills}
JOB DATA: {job_data}
Output PURE JSON matching this exact schema — NOTHING ELSE:
{{
  "personal": {{ ... }}, // copy verbatim + optional tailored title
  "summary": "3–5 sentence tailored summary",
  "flagship_projects": [ {{ "name": "str", "tech": ["str"], "description": "rephrased str" }}, ... ],
  "experience": [ {{ "company": "str", "title": "str", "start_date": "str", "end_date": "str", "bullets": ["bullet1", ...] }}, ... ],
  "skills": ["PrioritizedSkill1", ...],
  "education": [ {{ "degree": "str", "institution": "str", "location": "str" }}, ... ]
}}
"""

def load_master():
    with open(MASTER_ROOT / "master_career_data.json", "r", encoding="utf-8") as f:
        career = json.load(f)
    with open(MASTER_ROOT / "skills.json", "r", encoding="utf-8") as f:
        skills = json.load(f)
    return career, skills


def load_tailored_job(uuid_str: str):
    job_dir = JOB_ROOT / uuid_str
    tailored_dir = job_dir / "tailored"
    files = list(tailored_dir.glob("tailored_data_*.yaml"))
    if not files:
        raise FileNotFoundError(f"No tailored YAML found in {tailored_dir}")
    latest = max(files, key=lambda p: p.stat().st_mtime)
    with open(latest, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def call_llm(prompt_text: str, llm: str = "grok"):
    if GrokClient is None:
        print("\n" + "="*60)
        print(" MOCK / MANUAL MODE (GrokClient not available)")
        print(" 1. Copy the prompt below and paste it into Grok")
        print(" 2. Get the full JSON response from Grok")
        print(" 3. Paste the JSON back here and press Enter\n")
        print("="*60 + "\n")
        print(prompt_text[:1500] + "\n... (truncated - copy full prompt from console or file) ...\n")
        user_input = input("Paste the complete JSON output from Grok here:\n")
        try:
            return json.loads(user_input.strip())
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON pasted:\n{user_input[:200]}\n...") from e

    client = GrokClient()
    response = client.query(
        prompt_text,
        model="grok-beta",
        temperature=0.0,
        max_tokens=4500
    )
    try:
        return json.loads(response.strip())
    except json.JSONDecodeError as e:
        raise ValueError(f"LLM did not return valid JSON:\n{response[:400]}\n...") from e


def load_existing_json(out_dir: Path, from_json: str = None):
    if from_json:
        json_path = Path(from_json)
        if not json_path.is_file():
            raise FileNotFoundError(f"Specified JSON not found: {json_path}")
    else:
        files = list(out_dir.glob("resume_structured_*.json"))
        if not files:
            raise FileNotFoundError(f"No existing structured JSON found in {out_dir}. Run without --skip-llm first.")
        json_path = max(files, key=lambda p: p.stat().st_mtime)
        print(f"Using latest JSON: {json_path}")
    with open(json_path, "r", encoding="utf-8") as f:
        return json.load(f)


def normalize_url(url: str) -> str:
    """Ensure URL has protocol; fallback to https if missing."""
    url = url.strip()
    if not url:
        return ""
    if not url.startswith(('http://', 'https://')):
        return 'https://' + url
    return url


def render_markdown(tailored: dict) -> str:
    lines = []
    p = tailored["personal"]
    lines.append(f"# {p.get('name', 'Sean Luka Girgis')}")
    lines.append(f"**{p.get('title', '')}** | {p.get('location', '')}")
    
    # Use portfolio_url instead of website
    portfolio_url = normalize_url(p.get('portfolio_url', ''))
    lines.append(
        f"{p.get('phone')} • {p.get('email')} • "
        f"[LinkedIn]({p.get('linkedin', '')}) • "
        f"[GitHub]({p.get('github', '')}) • "
        f"[Portfolio]({portfolio_url})"
    )
    
    lines.append("")
    lines.append("## Professional Summary")
    lines.append(tailored["summary"])
    lines.append("")
    lines.append("## Flagship Projects")
    for proj in tailored["flagship_projects"]:
        lines.append(f"### {proj['name']}")
        lines.append(f"**Tech:** {', '.join(proj['tech'])}")
        lines.append(proj["description"])
        lines.append("")
    lines.append("## Professional Experience")
    for exp in tailored["experience"]:
        lines.append(f"**{exp['title']}** — **{exp['company']}**")
        lines.append(f"{exp['start_date']} – {exp['end_date']}")
        for b in exp["bullets"]:
            lines.append(f"- {b}")
        lines.append("")
    lines.append("## Skills")
    lines.append(", ".join(tailored["skills"]))
    lines.append("")
    lines.append("## Education")
    for edu in tailored["education"]:
        lines.append(f"- **{edu['degree']}** — {edu['institution']}, {edu['location']}")
    return "\n".join(lines)


def add_hyperlink(paragraph, text, url):
    """
    Add a real clickable hyperlink to a paragraph using low-level XML.
    """
    url = normalize_url(url)
    if not url:
        paragraph.add_run(text)
        return None

    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')

    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')  # blue
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(color)
    rPr.append(u)
    new_run.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return None


def render_docx(tailored: dict, output_path: Path):
    doc = Document()

    # Nicer margins
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # Header / Contact
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    r = p.add_run(f"{tailored['personal'].get('name', 'Sean Luka Girgis')}\n")
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "Calibri"

    r = p.add_run(f"{tailored['personal'].get('title', '')}\n")
    r.italic = True
    r.font.size = Pt(12)
    r.font.name = "Calibri"

    p.add_run(f"{tailored['personal'].get('phone')} • {tailored['personal'].get('email')} • ")

    add_hyperlink(p, "LinkedIn", tailored['personal'].get('linkedin', ''))
    p.add_run(" • ")

    add_hyperlink(p, "GitHub", tailored['personal'].get('github', ''))
    p.add_run(" • ")

    # Use portfolio_url instead of website
    add_hyperlink(p, "Portfolio", tailored['personal'].get('portfolio_url', ''))

    doc.add_paragraph()  # spacer

    # Summary
    h = doc.add_heading("Professional Summary", level=1)
    h.runs[0].font.name = "Calibri"
    h.runs[0].font.size = Pt(14)
    doc.add_paragraph(tailored["summary"])
    doc.add_paragraph()  # spacer

    # Flagship Projects
    h = doc.add_heading("Flagship Projects", level=1)
    h.runs[0].font.name = "Calibri"
    h.runs[0].font.size = Pt(14)

    for proj in tailored["flagship_projects"]:
        h = doc.add_heading(proj["name"], level=2)
        h.runs[0].font.name = "Calibri"

        p = doc.add_paragraph()
        p.add_run("Tech: ").bold = True
        p.add_run(", ".join(proj["tech"]))

        doc.add_paragraph(proj["description"])
        doc.add_paragraph()  # spacer between projects

    # Experience
    h = doc.add_heading("Professional Experience", level=1)
    h.runs[0].font.name = "Calibri"
    h.runs[0].font.size = Pt(14)

    for exp in tailored["experience"]:
        p = doc.add_paragraph()
        p.add_run(f"{exp['title']} — {exp['company']}").bold = True
        p.add_run(f" {exp['start_date']} – {exp['end_date']}").italic = True

        for bullet in exp["bullets"]:
            doc.add_paragraph(bullet, style="List Bullet")

        doc.add_paragraph()  # spacer

    # Skills
    h = doc.add_heading("Skills", level=1)
    h.runs[0].font.name = "Calibri"
    h.runs[0].font.size = Pt(14)
    doc.add_paragraph(", ".join(tailored["skills"]))
    doc.add_paragraph()  # spacer

    # Education
    h = doc.add_heading("Education", level=1)
    h.runs[0].font.name = "Calibri"
    h.runs[0].font.size = Pt(14)

    for edu in tailored["education"]:
        p = doc.add_paragraph()
        p.add_run(f"{edu['degree']}").bold = True
        p.add_run(f" — {edu['institution']}, {edu['location']}")

    # Apply consistent font fallback
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.name is None:
                run.font.name = "Calibri"
            if run.font.size is None or run.font.size.pt < 11:
                run.font.size = Pt(11)

    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(description="Generate tailored resume")
    parser.add_argument("--uuid", required=True, help="Job UUID")
    parser.add_argument("--version", default="v1", help="Output version tag")
    parser.add_argument("--llm", default="grok", help="LLM backend (grok default)")
    parser.add_argument("--skip-llm", action="store_true", help="Skip LLM call and re-render from existing JSON")
    parser.add_argument("--from-json", type=str, default=None, help="Path to existing JSON to re-render (with --skip-llm)")
    args = parser.parse_args()

    out_dir = JOB_ROOT / args.uuid / "generated"
    out_dir.mkdir(parents=True, exist_ok=True)

    if args.skip_llm:
        print("Skipping LLM — loading existing JSON for re-render...")
        tailored = load_existing_json(out_dir, args.from_json)
    else:
        # Load data
        master_career, master_skills = load_master()
        job_data = load_tailored_job(args.uuid)

        # Build prompt
        prompt_text = RESUME_PROMPT.format(
            master_career_data=json.dumps(master_career, indent=2),
            master_skills=json.dumps(master_skills, indent=2),
            job_data=yaml.dump(job_data, sort_keys=False, allow_unicode=True)
        )

        print("Generating resume content...")
        try:
            tailored = call_llm(prompt_text, args.llm)
        except Exception as e:
            print(f"LLM call failed: {e}")
            raise

    # Save outputs
    json_path = out_dir / f"resume_structured_{args.version}.json"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(tailored, f, indent=2, ensure_ascii=False)
    print(f"Saved structured JSON → {json_path}")

    md_content = render_markdown(tailored)
    md_path = out_dir / f"resume_preview_{args.version}.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_content)
    print(f"Saved Markdown preview → {md_path}")

    docx_path = out_dir / f"resume_{args.version}.docx"
    render_docx(tailored, docx_path)
    print(f"Saved DOCX → {docx_path}")


if __name__ == "__main__":
    main()