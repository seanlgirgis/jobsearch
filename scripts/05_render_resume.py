#!/usr/bin/env python3
"""
scripts/05_render_resume.py

Phase 5 in POC pipeline:
After intermediate JSON is generated (phase 04) → render to Markdown preview and DOCX.
Supports trimming for 1-2 page resumes (up to 5 recent jobs + summary block).
Includes --all to generate both trimmed and untrimmed versions at once.

Dependencies: pip install python-docx

Usage:
    python scripts/05_render_resume.py --uuid <uuid-or-prefix> --version tailored-v1 [--trim | --all]
"""

import argparse
import json
from pathlib import Path
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

JOB_ROOT = Path("data/jobs")
MASTER_ROOT = Path("data/master")

def resolve_job_folder(uuid_str: str) -> Path:
    """Smart resolution consistent with phases 02-04"""
    uuid_str = uuid_str.strip()
    job_dir = JOB_ROOT / uuid_str

    if job_dir.is_dir():
        return job_dir

    candidates = list(JOB_ROOT.glob(f"*_{uuid_str}*"))
    if not candidates:
        short = uuid_str[:8]
        candidates = list(JOB_ROOT.glob(f"*_{short}*"))

    if len(candidates) == 1:
        print(f"Resolved UUID/prefix → using folder: {candidates[0].name}")
        return candidates[0]
    elif len(candidates) > 1:
        print("Ambiguous UUID — multiple matches:")
        for c in candidates:
            print(f"  - {c.name}")
        raise ValueError(f"Ambiguous UUID: {uuid_str}")
    else:
        raise ValueError(
            f"Job folder not found for '{uuid_str}'\n"
            f"Expected: data/jobs/{uuid_str} or data/jobs/*_{uuid_str}*"
        )


def load_latest_intermediate(job_dir: Path, version: str) -> Dict[str, Any]:
    generated_dir = job_dir / "generated"
    if not generated_dir.is_dir():
        raise FileNotFoundError(f"No generated/ folder: {generated_dir}")

    # Look for exact match first
    exact_path = generated_dir / f"resume_intermediate_{version}.json"
    if exact_path.is_file():
        print(f"Using exact intermediate file: {exact_path.name}")
        with open(exact_path, "r", encoding="utf-8") as f:
            return json.load(f)

    # Fallback to latest if not found
    files = sorted(generated_dir.glob("resume_intermediate_*.json"), reverse=True)
    if not files:
        raise FileNotFoundError(f"No resume_intermediate_*.json found in {generated_dir}")

    latest = files[0]
    print(f"Exact version not found — using latest intermediate: {latest.name}")
    with open(latest, "r", encoding="utf-8") as f:
        return json.load(f)


def load_exclusions() -> set:
    career_path = MASTER_ROOT / "master_career_data.json"
    if not career_path.is_file():
        print("Warning: master_career_data.json not found — no exclusions applied")
        return set()
    with open(career_path, "r", encoding="utf-8") as f:
        master = json.load(f)
    excluded = {exp.get("company", "").lower() for exp in master.get("experience", []) if exp.get("exclude_from_resume", False)}
    if excluded:
        print(f"Applying exclusions for companies: {', '.join(excluded)}")
    return excluded


def clean_url_for_display(url: str) -> str:
    if not url:
        return ""
    cleaned = url.replace("https://", "").replace("http://", "").replace("www.", "").rstrip("/")
    return cleaned or url


def render_markdown(tailored: Dict, trim: bool = False, exclusions: set = set()) -> str:
    md = "# Tailored Resume Preview\n\n"

    personal = tailored.get("personal", {})
    name = personal.get("full_name", "Name Missing")
    title = personal.get("preferred_title", "")
    phone = personal.get("phone", "")
    email = personal.get("email", "")
    linkedin = personal.get("linkedin", "")
    github = personal.get("github", "")
    website = personal.get("personal_website", "") or personal.get("portfolio", "")

    md += f"**{name}**  \n"
    if title:
        md += f"{title}  \n"
    md += f"{phone} | {email}  \n"

    links = []
    if linkedin:
        links.append(f"[LinkedIn]({linkedin})")
    if github:
        links.append(f"[GitHub]({github})")
    if website:
        links.append(f"[Portfolio]({website})")
    if links:
        md += " | ".join(links) + "  \n\n"

    # Summary
    md += "## Professional Summary\n"
    md += tailored.get("summary", "Summary missing") + "\n\n"

    # Experience (apply exclusions)
    experiences = [exp for exp in tailored.get("experience", []) if exp.get("company", "").lower() not in exclusions]
    md += "## Professional Experience\n"
    display_exps = experiences[:5] if trim else experiences
    for exp in display_exps:
        md += f"### {exp.get('title', 'Title')} at {exp.get('company', 'Company')}\n"
        md += f"{exp.get('start_date', '?')} – {exp.get('end_date', '?')}\n\n"
        for bullet in exp.get("bullets", [])[:7]:
            md += f"- {bullet}\n"
        md += "\n"

    if not trim and len(experiences) > 5:
        md += "## Additional Experience\n"
        for exp in experiences[5:]:
            md += f"- **{exp.get('title', 'Title')}** at {exp.get('company', 'Company')} ({exp.get('start_date', '?')} – {exp.get('end_date', '?')})\n"

    # Education
    md += "## Education\n"
    for edu in tailored.get("education", []):
        md += f"- **{edu.get('degree', '')}**, {edu.get('institution', '')}"
        if edu.get('location'):
            md += f", {edu.get('location')}"
        dates = edu.get("dates", "")
        if dates and dates.lower() not in ["not specified", "(not specified)", "", "unknown"]:
            md += f" ({dates})"
        md += "\n"

    # Skills
    md += "## Skills\n"
    skills = tailored.get("skills", [])
    skill_str = ", ".join(s.get("name", s) if isinstance(s, dict) else s for s in skills)
    md += skill_str + "\n\n"

    # Flagship Projects
    projects = tailored.get("projects", []) or tailored.get("flagship_projects", [])
    if projects:
        md += "## Flagship Projects\n"
        for proj in projects:
            md += f"### {proj.get('name', 'Project')}\n"
            md += f"{proj.get('description', '')}\n"
            if proj.get("technologies"):
                md += f"Technologies: {', '.join(proj['technologies'])}\n"
            repo = proj.get("repo", "")
            if repo:
                repo_display = clean_url_for_display(repo)
                if repo_display:
                    md += f"Repo: [{repo_display}]({repo})\n"
            md += "\n"

    return md


def add_hyperlink(paragraph, url: str, text: str):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    new_run.append(rPr)

    text_run = OxmlElement('w:t')
    text_run.text = text
    new_run.append(text_run)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def render_docx(tailored: Dict, out_path: Path, trim: bool = False, exclusions: set = set()):
    doc = Document()

    personal = tailored.get("personal", {})

    # === Header - compact ===
    header_p = doc.add_paragraph()
    header_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_p.paragraph_format.space_after = Pt(0)  # ← reduce space after this paragraph

    # Name - 16 pt
    run_name = header_p.add_run(personal.get("full_name", "Name Missing"))
    run_name.bold = True
    run_name.font.size = Pt(16)

    # Phone + Email - next line in same paragraph
    header_p.add_run("\n")
    run_contact = header_p.add_run(f"{personal.get('phone', '')} | {personal.get('email', '')}")
    run_contact.font.size = Pt(11)

    # Links row - separate paragraph, no extra space
    links_p = doc.add_paragraph()
    links_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    links_p.paragraph_format.space_before = Pt(0)  # ← reduce space before links
    links_p.paragraph_format.space_after = Pt(6)  # ← small space after links

    linkedin = personal.get("linkedin", "")
    github = personal.get("github", "")
    website = personal.get("personal_website", "") or personal.get("portfolio", "")

    first = True
    if linkedin:
        if not first: links_p.add_run(" | ")
        add_hyperlink(links_p, linkedin, "LinkedIn")
        first = False
    if github:
        if not first: links_p.add_run(" | ")
        add_hyperlink(links_p, github, "GitHub")
        first = False
    if website:
        if not first: links_p.add_run(" | ")
        add_hyperlink(links_p, website, "Portfolio")
        first = False

    # Summary section - add a bit of space before first heading
    summary_heading = doc.add_heading("Professional Summary", level=1)
    summary_heading.paragraph_format.space_before = Pt(12)  # ← slight space before summary
    doc.add_paragraph(tailored.get("summary", "Summary missing"))

    # Experience
    doc.add_heading("Professional Experience", level=1)
    experiences = [exp for exp in tailored.get("experience", []) if exp.get("company", "").lower() not in exclusions]
    display_exps = experiences[:5] if trim else experiences
    for exp in display_exps:
        p = doc.add_paragraph()
        run = p.add_run(f"{exp.get('title', 'Title')} at {exp.get('company', 'Company')}")
        run.bold = True
        p.add_run(f"\n{exp.get('start_date', '?')} – {exp.get('end_date', '?')}")

        for bullet in exp.get("bullets", [])[:7]:
            doc.add_paragraph(bullet, style="ListBullet")

    if not trim and len(experiences) > 5:
        doc.add_heading("Additional Experience", level=1)
        for exp in experiences[5:]:
            p = doc.add_paragraph()
            run = p.add_run(f"{exp.get('title', 'Title')} at {exp.get('company', 'Company')}")
            run.bold = True
            p.add_run(f" ({exp.get('start_date', '?')} – {exp.get('end_date', '?')})")

    # Education
    doc.add_heading("Education", level=1)
    for edu in tailored.get("education", []):
        p = doc.add_paragraph()
        run = p.add_run(edu.get("degree", ""))
        run.bold = True
        line = f"\n{edu.get('institution', '')}"
        if edu.get('location'):
            line += f", {edu.get('location')}"
        dates = edu.get("dates", "")
        if dates and dates.lower() not in ["not specified", "(not specified)", "", "unknown"]:
            line += f" ({dates})"
        p.add_run(line)

    # Skills
    doc.add_heading("Skills", level=1)
    skills = tailored.get("skills", [])
    skill_str = ", ".join(s.get("name", s) if isinstance(s, dict) else s for s in skills)
    doc.add_paragraph(skill_str)

    # Flagship Projects
    projects = tailored.get("projects", []) or tailored.get("flagship_projects", [])
    if projects:
        doc.add_heading("Flagship Projects", level=1)
        for proj in projects:
            p = doc.add_paragraph()
            run = p.add_run(proj.get("name", "Project"))
            run.bold = True
            # Description - split sentences
            desc = proj.get('description', '')
            if desc:
                # heuristic: split on ". " to new lines
                desc_formatted = desc.replace(". ", ".\n")
                p.add_run(f"\n{desc_formatted}")

            if proj.get("technologies"):
                p.add_run("\n")
                run_tech = p.add_run("Technologies:")
                run_tech.bold = True
                p.add_run(f" {', '.join(proj['technologies'])}")

            repo = proj.get("repo", "")
            if repo:
                repo_display = clean_url_for_display(repo)
                if repo_display:
                    p.add_run("\n")
                    run_repo = p.add_run("Repo: ")
                    run_repo.bold = True
                    add_hyperlink(p, repo, repo_display)

    # Formatting
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial"
            if paragraph.style.name.startswith("Heading"):
                run.font.size = Pt(13)
            elif run.font.size == Pt(16):
                pass
            else:
                run.font.size = Pt(11)

    doc.save(out_path)


def generate_files(tailored: Dict, job_folder: Path, version: str, trim_mode: bool, exclusions: set):
    suffix = "_trimmed" if trim_mode else ""
    md_path = job_folder / "generated" / f"resume_preview_{version}{suffix}.md"
    docx_path = job_folder / "generated" / f"resume_{version}{suffix}.docx"

    with open(md_path, "w", encoding="utf-8") as f:
        f.write(render_markdown(tailored, trim=trim_mode, exclusions=exclusions))
    print(f"Markdown saved → {md_path}")

    render_docx(tailored, docx_path, trim=trim_mode, exclusions=exclusions)
    print(f"DOCX saved → {docx_path}")


def main():
    parser = argparse.ArgumentParser(description="Phase 5: Render resume from intermediate JSON to MD & DOCX")
    parser.add_argument("--uuid", required=True, help="Job UUID or short prefix (e.g. cdb9a3fa)")
    parser.add_argument("--version", default="v1", help="Intermediate JSON version tag (e.g. tailored-v1)")
    group = parser.add_mutually_exclusive_group()
    group.add_argument("--trim", action="store_true", help="Generate trimmed version only (5 recent jobs)")
    group.add_argument("--all", action="store_true", help="Generate both trimmed and full versions")
    args = parser.parse_args()

    job_folder = resolve_job_folder(args.uuid)
    print(f"Processing job: {job_folder.name}")

    tailored = load_latest_intermediate(job_folder, args.version)
    exclusions = load_exclusions()

    if args.all:
        print("Generating both trimmed and full versions...")
        generate_files(tailored, job_folder, args.version, trim_mode=True, exclusions=exclusions)
        generate_files(tailored, job_folder, args.version, trim_mode=False, exclusions=exclusions)
    else:
        trim_mode = args.trim
        mode_name = "trimmed" if trim_mode else "full"
        print(f"Generating {mode_name} version...")
        generate_files(tailored, job_folder, args.version, trim_mode=trim_mode, exclusions=exclusions)

    print("\nDone.")


if __name__ == "__main__":
    main()